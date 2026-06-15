from fastapi import APIRouter, UploadFile, File
from core.db import add_document, get_documents, delete_document, get_conn
from core.vector_store import chunk_text, add_document_vectors, delete_document_vectors, is_model_ready
import os

router = APIRouter(prefix="/knowledge", tags=["knowledge"])

UPLOAD_DIR = os.path.join("data", "knowledge")
os.makedirs(UPLOAD_DIR, exist_ok=True)

ALLOWED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx", ".html", ".htm"}

# 待补建索引队列（模型未加载时暂存 doc_id + 已解析文本）
_pending_index: list = []  # [(doc_id, text), ...]


def _read_file(filepath: str, filename: str) -> str:
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".pdf":
        from pypdf import PdfReader
        reader = PdfReader(filepath)
        return "\n\n".join(page.extract_text() or "" for page in reader.pages)
    elif ext == ".docx":
        from docx import Document
        doc = Document(filepath)
        parts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            if para.style and para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading", "").strip()
                try:
                    level = int(level)
                except ValueError:
                    level = 1
                parts.append(f"{'#' * level} {text}")
            else:
                parts.append(text)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n\n".join(parts)
    elif ext in (".html", ".htm"):
        from bs4 import BeautifulSoup
        try:
            with open(filepath, "r", encoding="utf-8") as f:
                raw = f.read()
        except UnicodeDecodeError:
            with open(filepath, "r", encoding="gbk", errors="replace") as f:
                raw = f.read()
        soup = BeautifulSoup(raw, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        parts = []
        for tag in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "td", "th", "article", "section"]):
            text = tag.get_text(strip=True)
            if not text or len(text) < 2:
                continue
            if tag.name.startswith("h"):
                level = int(tag.name[1])
                parts.append(f"{'#' * level} {text}")
            else:
                parts.append(text)
        return "\n\n".join(parts) if parts else soup.get_text(separator="\n", strip=True)
    else:
        with open(filepath, "r", encoding="utf-8") as f:
            return f.read()


def _flush_pending_index():
    if not _pending_index or not is_model_ready():
        return
    doc_ids = [item[0] for item in _pending_index]
    with get_conn() as conn:
        cursor = conn.cursor()
        placeholders = ",".join("?" * len(doc_ids))
        cursor.execute(f"SELECT id, filename FROM documents WHERE id IN ({placeholders})", doc_ids)
        id_to_file = {r["id"]: r["filename"] for r in cursor.fetchall()}
    flushed = []
    remaining = []
    for doc_id, text in _pending_index:
        try:
            filename = id_to_file.get(doc_id)
            if not filename:
                continue
            chunks = chunk_text(text)
            count = add_document_vectors(doc_id, chunks, filename)
            flushed.append(f"{filename}({count}条)")
        except Exception as e:
            remaining.append((doc_id, text))
            print(f"[知识库] 补建索引失败 doc_id={doc_id}: {e}")
    _pending_index.clear()
    _pending_index.extend(remaining)
    if flushed:
        print(f"[知识库] 补建索引完成: {', '.join(flushed)}")


@router.post("/upload")
async def upload_knowledge(file: UploadFile = File(...)):
    if not file.filename:
        return {"status": "error", "message": "未选择文件"}
    safe_name = os.path.basename(file.filename)
    if not safe_name or safe_name.startswith('.') or safe_name != file.filename:
        return {"status": "error", "message": "文件名不合法"}
    ext = os.path.splitext(safe_name)[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        return {"status": "error", "message": f"不支持的文件类型: {ext}，仅支持 {', '.join(ALLOWED_EXTENSIONS)}"}

    filepath = os.path.join(UPLOAD_DIR, safe_name)
    total_bytes = 0
    MAX_SIZE = 50 * 1024 * 1024  # 50MB
    try:
        with open(filepath, "wb") as f:
            while chunk := await file.read(1024 * 1024):
                total_bytes += len(chunk)
                if total_bytes > MAX_SIZE:
                    os.remove(filepath)
                    return {"status": "error", "message": "文件大小不能超过50MB"}
                f.write(chunk)
    except Exception as e:
        return {"status": "error", "message": f"文件保存失败: {str(e)}"}

    try:
        text = _read_file(filepath, file.filename)
    except Exception as e:
        os.remove(filepath)
        return {"status": "error", "message": f"文件读取失败: {str(e)}"}

    if not text or not text.strip():
        os.remove(filepath)
        return {"status": "error", "message": "文件内容为空"}

    size = os.path.getsize(filepath)
    chunks = chunk_text(text)
    doc_id = add_document(file.filename, len(chunks), size)

    if is_model_ready():
        count = add_document_vectors(doc_id, chunks, file.filename)
    else:
        count = 0
        _pending_index.append((doc_id, text[:200000] if len(text) > 200000 else text))

    return {"status": "ok", "id": doc_id, "filename": file.filename, "chunks": len(chunks), "indexed": count}


@router.get("/")
async def list_knowledge():
    _flush_pending_index()
    docs = get_documents()
    return {"documents": docs, "model_ready": is_model_ready()}


@router.delete("/{doc_id}")
async def remove_knowledge(doc_id: int):
    _pending_index[:] = [(did, txt) for did, txt in _pending_index if did != doc_id]
    # 先查文件名再删记录
    with get_conn() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT filename FROM documents WHERE id = ?", (doc_id,))
        row = cursor.fetchone()
    if row:
        filepath = os.path.join(UPLOAD_DIR, row["filename"])
        try:
            os.remove(filepath)
        except Exception as e:
            print(f"[知识库] 删除文件失败: {e}")
    delete_document(doc_id)
    delete_document_vectors(doc_id)
    return {"status": "ok"}
