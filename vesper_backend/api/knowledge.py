"""知识库管理 API —— 文档上传、解析、向量化索引与检索（per-character）。"""
import logging
import os

from fastapi import APIRouter, UploadFile, File, Query
from core.db import add_document, get_documents, delete_document, get_chat_conn
from core.vector_store import chunk_text, add_document_vectors, delete_document_vectors, is_model_ready

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/knowledge", tags=["knowledge"])

UPLOAD_DIR = os.path.join("data", "knowledge")
os.makedirs(UPLOAD_DIR, exist_ok=True)

ALLOWED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx", ".html", ".htm"}

# 待补建索引队列（模型未加载时暂存 doc_id + 已解析文本）
_pending_index: list = []  # [(doc_id, text), ...]


def _read_file(filepath: str, filename: str) -> str:
    """根据文件扩展名选择解析方式，提取文本内容（支持 PDF/DOCX/HTML/TXT）。"""
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".pdf":
        from pypdf import PdfReader
        reader = PdfReader(filepath)
        return "\n\n".join(page.extract_text() or "" for page in reader.pages)
    elif ext == ".docx":
        from docx import Document
        doc = Document(filepath)
        parts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            if para.style and para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading", "").strip()
                try:
                    level = int(level)
                except ValueError:
                    level = 1
                parts.append(f"{'#' * level} {text}")
            else:
                parts.append(text)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n\n".join(parts)
    elif ext in (".html", ".htm"):
        from bs4 import BeautifulSoup
        try:
            with open(filepath, "r", encoding="utf-8") as f:
                raw = f.read()
        except UnicodeDecodeError:
            with open(filepath, "r", encoding="gbk", errors="replace") as f:
                raw = f.read()
        soup = BeautifulSoup(raw, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        parts = []
        for tag in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "td", "th", "article", "section"]):
            text = tag.get_text(strip=True)
            if not text or len(text) < 2:
                continue
            if tag.name.startswith("h") and len(tag.name) > 1:
                try:
                    level = int(tag.name[1])
                except (ValueError, IndexError):
                    continue
                parts.append(f"{'#' * level} {text}")
            else:
                parts.append(text)
        return "\n\n".join(parts) if parts else soup.get_text(separator="\n", strip=True)
    else:
        with open(filepath, "r", encoding="utf-8") as f:
            return f.read()


def _flush_pending_index():
    """将待建索引队列中的文档批量写入向量库（模型就绪时触发）。"""
    if not _pending_index or not is_model_ready():
        return
    doc_ids = [item[0] for item in _pending_index]
    # 从所有角色库查找文档名（_pending_index 存储时已带 character_id）
    id_to_file = {}
    for cid in range(10):  # 查前 10 个角色
        try:
            with get_chat_conn(cid) as conn:
                cursor = conn.cursor()
                placeholders = ",".join("?" * len(doc_ids))
                cursor.execute(f"SELECT id, filename FROM documents WHERE id IN ({placeholders})", doc_ids)
                for r in cursor.fetchall():
                    id_to_file[r["id"]] = r["filename"]
            if len(id_to_file) == len(doc_ids):
                break
        except Exception:
            pass
    flushed = []
    remaining = []
    for doc_id, text in _pending_index:
        try:
            filename = id_to_file.get(doc_id)
            if not filename:
                continue
            chunks = chunk_text(text)
            count = add_document_vectors(doc_id, chunks, filename)
            flushed.append(f"{filename}({count}条)")
        except Exception as e:
            remaining.append((doc_id, text))
            logger.warning(f"[知识库] 补建索引失败 doc_id={doc_id}: {e}")
    _pending_index.clear()
    _pending_index.extend(remaining)
    if flushed:
        logger.info(f"[知识库] 补建索引完成: {', '.join(flushed)}")


@router.post("/upload")
async def upload_knowledge(file: UploadFile = File(...),
                          character_id: int = Query(0, description="角色 ID，默认 0")):
    """上传知识库文档到指定角色（per-character）。"""
    if not file.filename:
        return {"status": "error", "message": "未选择文件"}
    safe_name = os.path.basename(file.filename)
    if not safe_name or safe_name.startswith('.') or safe_name != file.filename:
        return {"status": "error", "message": "文件名不合法"}
    ext = os.path.splitext(safe_name)[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        return {"status": "error", "message": f"不支持的文件类型: {ext}，仅支持 {', '.join(ALLOWED_EXTENSIONS)}"}

    filepath = os.path.join(UPLOAD_DIR, safe_name)
    total_bytes = 0
    MAX_SIZE = 50 * 1024 * 1024  # 50MB
    try:
        with open(filepath, "wb") as f:
            while chunk := await file.read(1024 * 1024):
                total_bytes += len(chunk)
                if total_bytes > MAX_SIZE:
                    os.remove(filepath)
                    return {"status": "error", "message": "文件大小不能超过50MB"}
                f.write(chunk)
    except Exception as e:
        return {"status": "error", "message": f"文件保存失败: {str(e)}"}

    try:
        text = _read_file(filepath, file.filename)
    except Exception as e:
        os.remove(filepath)
        return {"status": "error", "message": f"文件读取失败: {str(e)}"}

    if not text or not text.strip():
        os.remove(filepath)
        return {"status": "error", "message": "文件内容为空"}

    size = os.path.getsize(filepath)
    chunks = chunk_text(text)
    doc_id = add_document(file.filename, len(chunks), size, character_id=character_id)

    if is_model_ready():
        count = add_document_vectors(doc_id, chunks, file.filename, character_id=character_id)
    else:
        count = 0
        _pending_index.append((doc_id, text[:200000] if len(text) > 200000 else text))

    return {"status": "ok", "id": doc_id, "filename": file.filename, "chunks": len(chunks), "indexed": count}


@router.get("/")
async def list_knowledge(character_id: int = Query(0, description="角色 ID，默认 0")):
    """列出指定角色的知识库文档。"""
    _flush_pending_index()
    docs = get_documents(character_id=character_id)
    return {"documents": docs, "model_ready": is_model_ready()}


@router.delete("/{doc_id}")
async def remove_knowledge(doc_id: int,
                          character_id: int = Query(0, description="角色 ID，默认 0")):
    """删除指定角色知识库文档及其向量索引。"""
    _pending_index[:] = [(did, txt) for did, txt in _pending_index if did != doc_id]
    with get_chat_conn(character_id) as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT filename FROM documents WHERE id = ?", (doc_id,))
        row = cursor.fetchone()
    if row:
        filename = row["filename"]
        # 路径安全校验
        from core.security import validate_path_within
        filepath = os.path.join(UPLOAD_DIR, filename)
        if not validate_path_within(UPLOAD_DIR, filepath):
            raise HTTPException(400, "文件名包含非法路径")
        try:
            os.remove(filepath)
        except Exception as e:
            logger.warning(f"[知识库] 删除文件失败: {e}")
    delete_document(doc_id)
    delete_document_vectors(doc_id)
    return {"status": "ok"}
